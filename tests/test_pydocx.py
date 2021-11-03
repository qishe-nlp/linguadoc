from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from linguadoc.lib import read_json 
from datetime import datetime

def test_sample():
  jsonfile = "./en_doc.json"
  data = read_json(jsonfile)

  document = Document('demo_details.docx')
  document._body.clear_content()
  document.add_heading('Swan Lake', level=0)
  document.add_paragraph('歧舌', style='Quote')
  created_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
  document.add_paragraph(created_time, style='Quote')

  for s in data:
    document.add_heading("{}\n{}".format(s["sentence"], s["translation"]), level=1)

    document.add_heading("句型结构", level=2)

    document.add_heading(s["structure"], level=3)
    analysis = s["analysis"]
    for ele in analysis:
      if ele["explanation"]:
        p = document.add_paragraph("{}    {}".format(ele["text"], ele["translation"]), style='List Bullet')

    document.add_heading("知识要点", level=2)

    kg = s["kg"]
    for key, value in kg.items():
      for v in value: 
        p = document.add_paragraph("{}    {}".format(key, v["text"]), style='List Bullet')

  document.add_paragraph("This is green", style='test')
  document.add_paragraph("我是绿色的", style='test')

  document.save('result_full.docx')

