from docx import Document
from docx.shared import Cm
from linguadoc.lib import read_json
from datetime import datetime
import os

class LinguaDoc:
  """It is used for docx generation of sentence structure and KG
  """

  def __init__(self, jsonfile, lang):
    """Initialize docx template, source data
    Args:
      jsonfile (str): json file, whose content is written into doc 
      lang (str): language of the jsonfile
    """

    self.lang = lang
    _template_dir = os.path.dirname(__file__) + '/' + lang
    self.template = os.path.join(_template_dir, lang+'_template.docx')
    self.cover = os.path.join(_template_dir, lang+'_doc_cover.jpg')
    self.end = os.path.join(_template_dir, lang+'_doc_end.jpg')
    self.jsonfile = jsonfile
    self.data = read_json(self.jsonfile)

  def gen_doc(self, title, output="result.docx"):
    """Generate docx file
    Args:
      title (str): title written in ppt home slide
      output (str): output docx file name 
    """
    document = Document(self.template)
    document._body.clear_content()
    document.add_paragraph(title, style="Title")
    document.add_paragraph(title, style='Subtitle')
    created_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    document.add_paragraph('歧舌-'+created_time, style='Quote')
    document.add_picture(self.cover, width=Cm(15.24))
    document.add_page_break()

    for s in self.data:
      sentence = s["sentence"]
      document.add_heading("{}\n{}".format(sentence["text"], sentence["meaning"]), level=1)

      document.add_heading("句型结构", level=2)

      document.add_heading(s["structure_rep"], level=3)
      structure = s["structure"]
      for ele in structure:
        if ele["explanation"]:
          p = document.add_paragraph("{}    {}".format(ele["text"], ele["meaning"]), style='List Bullet')

      document.add_heading("知识要点", level=2)

      kg = s["kg"]
      for key, value in kg.items():
        for v in value: 
          p = document.add_paragraph("{}    {}".format(key, v["text"]), style='List Bullet')

    document.add_picture(self.end, width=Cm(15.24))

    document.save(output)

